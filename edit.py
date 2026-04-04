import tkinter as tk
from tkinter import ttk, filedialog, messagebox, colorchooser, font, simpledialog
import os
from datetime import datetime
import io
import docx
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pyperclip
import msoffcrypto

class ModernWordEditor:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("专业便携式文档编辑器")
        self.root.geometry("1400x900")
        self.root.configure(bg='#f0f0f0')
        
        # 文件状态
        self.current_file = None
        self.is_modified = False
        self.current_font_size = 12
        self.current_font_family = "微软雅黑"
        
        # 页面设置（默认边距：单位厘米）
        self.page_margins = {
            'top': 2.54,
            'bottom': 2.54,
            'left': 3.17,
            'right': 3.17
        }
        
        # 创建界面
        self.create_menu_bar()
        self.create_toolbar()
        self.create_quick_bar()
        self.create_ruler()
        self.create_main_area()
        self.create_footer_bar()  # 新增：底部开源地址栏
        
        # 绑定事件：点击窗口关闭按钮
        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)
        
        # 绑定快捷键
        self.bind_events()
        
        # 初始化新文档
        self.new_file()
        
    def create_menu_bar(self):
        """创建菜单栏"""
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)
        
        # 文件菜单
        file_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="文件", menu=file_menu)
        file_menu.add_command(label="新建", command=self.new_file, accelerator="Ctrl+N")
        file_menu.add_command(label="打开", command=self.open_file, accelerator="Ctrl+O")
        file_menu.add_command(label="保存", command=self.save_file, accelerator="Ctrl+S")
        file_menu.add_command(label="另存为", command=self.save_as_file, accelerator="Ctrl+Shift+S")
        file_menu.add_separator()
        file_menu.add_command(label="加密保存为DOCX", command=self.encrypt_and_save_dialog)
        file_menu.add_separator()
        file_menu.add_command(label="退出", command=self.on_closing, accelerator="Ctrl+Q")
        
        # 编辑菜单
        edit_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="编辑", menu=edit_menu)
        edit_menu.add_command(label="撤销", command=self.undo, accelerator="Ctrl+Z")
        edit_menu.add_command(label="重做", command=self.redo, accelerator="Ctrl+Y")
        edit_menu.add_separator()
        edit_menu.add_command(label="剪切", command=self.cut_text, accelerator="Ctrl+X")
        edit_menu.add_command(label="复制", command=self.copy_text, accelerator="Ctrl+C")
        edit_menu.add_command(label="粘贴", command=self.paste_text, accelerator="Ctrl+V")
        edit_menu.add_separator()
        edit_menu.add_command(label="全选", command=self.select_all, accelerator="Ctrl+A")
        edit_menu.add_command(label="查找替换", command=self.find_replace, accelerator="Ctrl+F")
        
        # 格式菜单
        format_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="格式", menu=format_menu)
        format_menu.add_separator()
        format_menu.add_command(label="项目符号", command=self.add_bullets)
        format_menu.add_command(label="编号", command=self.add_numbering)

    def setup_page_margins(self):
        """页边距设置窗口"""
        dialog = tk.Toplevel(self.root)
        dialog.title("页面设置")
        dialog.geometry("300x250")
        dialog.configure(bg='#f0f0f0')
        
        tk.Label(dialog, text="页边距设置 (单位: 厘米)", bg='#f0f0f0', font=("微软雅黑", 10, "bold")).pack(pady=10)
        
        frame = tk.Frame(dialog, bg='#f0f0f0')
        frame.pack(pady=5)
        
        # 网格布局输入框
        tk.Label(frame, text="上:", bg='#f0f0f0').grid(row=0, column=0, padx=5, pady=5)
        top_e = tk.Entry(frame, width=5)
        top_e.insert(0, str(self.page_margins['top']))
        top_e.grid(row=0, column=1, padx=5, pady=5)
        
        tk.Label(frame, text="下:", bg='#f0f0f0').grid(row=0, column=2, padx=5, pady=5)
        bottom_e = tk.Entry(frame, width=5)
        bottom_e.insert(0, str(self.page_margins['bottom']))
        bottom_e.grid(row=0, column=3, padx=5, pady=5)
        
        tk.Label(frame, text="左:", bg='#f0f0f0').grid(row=1, column=0, padx=5, pady=5)
        left_e = tk.Entry(frame, width=5)
        left_e.insert(0, str(self.page_margins['left']))
        left_e.grid(row=1, column=1, padx=5, pady=5)
        
        tk.Label(frame, text="右:", bg='#f0f0f0').grid(row=1, column=2, padx=5, pady=5)
        right_e = tk.Entry(frame, width=5)
        right_e.insert(0, str(self.page_margins['right']))
        right_e.grid(row=1, column=3, padx=5, pady=5)
        
        def apply_margins():
            try:
                self.page_margins['top'] = float(top_e.get())
                self.page_margins['bottom'] = float(bottom_e.get())
                self.page_margins['left'] = float(left_e.get())
                self.page_margins['right'] = float(right_e.get())
                messagebox.showinfo("成功", "页边距已设置，保存为DOCX时生效")
                dialog.destroy()
            except ValueError:
                messagebox.showerror("错误", "请输入有效的数字")
        
        tk.Button(dialog, text="确定", command=apply_margins, bg='#3498db', fg='white', padx=20).pack(pady=15)

    def create_quick_bar(self):
        """创建第二行工具栏（新增页边距按钮）"""
        self.quick_bar = tk.Frame(self.root, bg='#e6e6e6', relief=tk.FLAT, bd=1)
        self.quick_bar.pack(side=tk.TOP, fill=tk.X, padx=2, pady=0)
        
        tk.Label(self.quick_bar, text="快速插入: ", bg='#e6e6e6', font=("微软雅黑", 9)).pack(side=tk.LEFT, padx=5)
        self.create_quick_button("📅 日期", self.insert_date).pack(side=tk.LEFT, padx=1, pady=2)
        self.create_quick_button("⏰ 时间", self.insert_time).pack(side=tk.LEFT, padx=1, pady=2)
        self.create_quick_button("📊 表格", self.insert_table).pack(side=tk.LEFT, padx=1, pady=2)
        self.create_quick_button("📏 分隔线", self.insert_separator).pack(side=tk.LEFT, padx=1, pady=2)
        
        tk.Frame(self.quick_bar, width=2, bg='#cccccc').pack(side=tk.LEFT, padx=10, fill=tk.Y, pady=3)
        
        # 新增：页边距设置按钮
        self.create_quick_button("📐 页边距", self.setup_page_margins).pack(side=tk.LEFT, padx=1, pady=2)
        
        tk.Frame(self.quick_bar, width=2, bg='#cccccc').pack(side=tk.LEFT, padx=10, fill=tk.Y, pady=3)
        
        tk.Label(self.quick_bar, text="统计: ", bg='#e6e6e6', font=("微软雅黑", 9)).pack(side=tk.LEFT, padx=5)
        self.word_count_label = tk.Label(self.quick_bar, text="字数: 0", bg='#e6e6e6', font=("微软雅黑", 9))
        self.word_count_label.pack(side=tk.LEFT, padx=5)
        self.char_count_label = tk.Label(self.quick_bar, text="字符: 0", bg='#e6e6e6', font=("微软雅黑", 9))
        self.char_count_label.pack(side=tk.LEFT, padx=5)
        self.para_count_label = tk.Label(self.quick_bar, text="段落: 1", bg='#e6e6e6', font=("微软雅黑", 9))
        self.para_count_label.pack(side=tk.LEFT, padx=5)

        tk.Frame(self.quick_bar, width=2, bg='#cccccc').pack(side=tk.RIGHT, padx=10, fill=tk.Y, pady=3)
        self.encrypt_btn = tk.Button(self.quick_bar, text="🔐 加密保存", command=self.encrypt_and_save_dialog,
                                      bg='#27ae60', fg='white', font=("微软雅黑", 9, "bold"), relief=tk.FLAT, padx=10)
        self.encrypt_btn.pack(side=tk.RIGHT, padx=5, pady=2)

    def create_quick_button(self, text, command):
        btn = tk.Button(self.quick_bar, text=text, command=command, bg='#e6e6e6', 
                       relief=tk.FLAT, padx=5, pady=2, font=("微软雅黑", 8))
        btn.bind('<Enter>', lambda e: btn.config(bg='#d0d0d0'))
        btn.bind('<Leave>', lambda e: btn.config(bg='#e6e6e6'))
        return btn

    def create_toolbar(self):
        """创建工具栏"""
        self.toolbar = tk.Frame(self.root, bg='#ffffff', relief=tk.RAISED, bd=1)
        self.toolbar.pack(side=tk.TOP, fill=tk.X, padx=2, pady=2)
        
        self.create_tool_button("📄 新建", self.new_file).pack(side=tk.LEFT, padx=2)
        self.create_tool_button("📂 打开", self.open_file).pack(side=tk.LEFT, padx=2)
        self.create_tool_button("💾 保存", self.save_file).pack(side=tk.LEFT, padx=2)
        
        tk.Frame(self.toolbar, width=2, bg='#cccccc').pack(side=tk.LEFT, padx=5, fill=tk.Y, pady=2)
        
        self.create_tool_button("𝐁", self.bold_text, width=3).pack(side=tk.LEFT, padx=1)
        self.create_tool_button("𝑰", self.italic_text, width=3).pack(side=tk.LEFT, padx=1)
        self.create_tool_button("U", self.underline_text, width=3).pack(side=tk.LEFT, padx=1)
        
        tk.Frame(self.toolbar, width=2, bg='#cccccc').pack(side=tk.LEFT, padx=5, fill=tk.Y, pady=2)
        
        tk.Label(self.toolbar, text="字体:", bg='#ffffff').pack(side=tk.LEFT, padx=2)
        self.font_family_combo = ttk.Combobox(self.toolbar, values=list(tk.font.families())[:50], width=12)
        self.font_family_combo.set("微软雅黑")
        self.font_family_combo.pack(side=tk.LEFT, padx=2)
        self.font_family_combo.bind('<<ComboboxSelected>>', self.change_font_family)
        
        tk.Label(self.toolbar, text="字号:", bg='#ffffff').pack(side=tk.LEFT, padx=2)
        self.font_size_combo = ttk.Combobox(self.toolbar, values=[8,9,10,11,12,14,16,18,20,22,24,26,28,36,48,72], width=5)
        self.font_size_combo.set(12)
        self.font_size_combo.pack(side=tk.LEFT, padx=2)
        self.font_size_combo.bind('<<ComboboxSelected>>', self.change_font_size)
        
        tk.Frame(self.toolbar, width=2, bg='#cccccc').pack(side=tk.LEFT, padx=5, fill=tk.Y, pady=2)
        
        self.create_tool_button("🎨 颜色", self.choose_text_color).pack(side=tk.LEFT, padx=2)
        
        tk.Frame(self.toolbar, width=2, bg='#cccccc').pack(side=tk.LEFT, padx=5, fill=tk.Y, pady=2)
        
        self.create_tool_button("⬅️", lambda: self.set_alignment('left'), width=3).pack(side=tk.LEFT, padx=1)
        self.create_tool_button("🀫", lambda: self.set_alignment('center'), width=3).pack(side=tk.LEFT, padx=1)
        self.create_tool_button("➡️", lambda: self.set_alignment('right'), width=3).pack(side=tk.LEFT, padx=1)

    def create_tool_button(self, text, command, width=None):
        btn = tk.Button(self.toolbar, text=text, command=command, bg='#ffffff', 
                       relief=tk.FLAT, padx=5, pady=3, font=("微软雅黑", 9))
        if width:
            btn.config(width=width)
        btn.bind('<Enter>', lambda e: btn.config(relief=tk.RAISED, bg='#e0e0e0'))
        btn.bind('<Leave>', lambda e: btn.config(relief=tk.FLAT, bg='#ffffff'))
        return btn
    
    def create_ruler(self):
        self.ruler_frame = tk.Frame(self.root, bg='#e8e8e8', height=25)
        self.ruler_frame.pack(side=tk.TOP, fill=tk.X, padx=2)
        self.ruler_canvas = tk.Canvas(self.ruler_frame, height=25, bg='#e8e8e8', highlightthickness=0)
        self.ruler_canvas.pack(fill=tk.X)
        for i in range(0, 100):
            x = 50 + i * 8
            if i % 10 == 0:
                self.ruler_canvas.create_line(x, 15, x, 25, width=1)
                self.ruler_canvas.create_text(x, 10, text=str(i), font=("Arial", 8))
            elif i % 5 == 0:
                self.ruler_canvas.create_line(x, 18, x, 25, width=1)
            else:
                self.ruler_canvas.create_line(x, 20, x, 25, width=1)
    
    def create_main_area(self):
        self.main_container = tk.Frame(self.root, bg='#e8e8e8')
        # 关键修改：移除底部padding，消除大空白
        self.main_container.pack(fill=tk.BOTH, expand=True, padx=10, pady=(10, 0))
        
        self.page_canvas = tk.Canvas(self.main_container, bg='#e8e8e8', highlightthickness=0)
        self.page_canvas.pack(fill=tk.BOTH, expand=True)
        
        self.paper_frame = tk.Frame(self.page_canvas, bg='#ffffff', relief=tk.SUNKEN, bd=1)
        self.paper_window = self.page_canvas.create_window(0, 0, window=self.paper_frame, anchor='nw')
        
        self.text_area = tk.Text(self.paper_frame, wrap=tk.WORD, font=("微软雅黑", 12),
                                 undo=True, autoseparators=True, maxundo=100,
                                 padx=60, pady=50, bg='#ffffff', fg='#000000',
                                 selectbackground='#3399ff', selectforeground='#ffffff',
                                 relief=tk.FLAT, borderwidth=0)
        self.text_area.pack(fill=tk.BOTH, expand=True)
        
        self.text_area.tag_configure("bold", font=("微软雅黑", 12, "bold"))
        self.text_area.tag_configure("italic", font=("微软雅黑", 12, "italic"))
        self.text_area.tag_configure("underline", font=("微软雅黑", 12, "underline"))
        self.text_area.tag_configure("center", justify='center')
        self.text_area.tag_configure("right", justify='right')
        
        self.text_area.bind('<<Modified>>', self.on_text_modified)
        self.text_area.bind('<KeyRelease>', self.update_stats)
        self.text_area.bind('<ButtonRelease-1>', self.update_stats)
        
        self.scrollbar = tk.Scrollbar(self.main_container, orient=tk.VERTICAL, command=self.text_area.yview)
        self.scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.text_area.config(yscrollcommand=self.scrollbar.set)
        
        self.page_canvas.bind('<Configure>', self.on_canvas_configure)
        self.paper_frame.bind('<Configure>', self.on_frame_configure)
    
    def create_footer_bar(self):
        """创建底部极简开源地址栏（高度仅20像素）"""
        self.footer_bar = tk.Frame(self.root, bg='#f0f0f0', relief=tk.FLAT, bd=0, height=20)
        self.footer_bar.pack(side=tk.BOTTOM, fill=tk.X, padx=0, pady=0)
        
        footer_label = tk.Label(self.footer_bar, text="本工具开源地址：https://github.com/lynvortex/editor-portable", 
                               bg='#f0f0f0', fg='#666666', font=("微软雅黑", 8))
        footer_label.pack(side=tk.RIGHT, padx=10, pady=2)
    
    def on_canvas_configure(self, event):
        width = event.width - 20
        self.page_canvas.itemconfig(self.paper_window, width=width)
        self.page_canvas.config(scrollregion=self.page_canvas.bbox("all"))
    
    def on_frame_configure(self, event):
        self.page_canvas.config(scrollregion=self.page_canvas.bbox("all"))
    
    def bind_events(self):
        self.root.bind('<Control-n>', lambda e: self.new_file())
        self.root.bind('<Control-o>', lambda e: self.open_file())
        self.root.bind('<Control-s>', lambda e: self.save_file())
        self.root.bind('<Control-z>', lambda e: self.undo())
        self.root.bind('<Control-y>', lambda e: self.redo())
        self.root.bind('<Control-a>', lambda e: self.select_all())
        self.root.bind('<Control-f>', lambda e: self.find_replace())
        self.root.bind('<Control-q>', lambda e: self.on_closing())
    
    # ================= 核心功能实现 =================
    
    def encrypt_and_save_dialog(self):
        """优化后的加密对话框：输入密码后直接保存"""
        dialog = tk.Toplevel(self.root)
        dialog.title("加密保存为 DOCX")
        dialog.geometry("400x260")
        dialog.configure(bg='#f0f0f0')
        dialog.resizable(False, False)
        
        tk.Label(dialog, text="设置文档打开密码", font=("微软雅黑", 12, "bold"), bg='#f0f0f0').pack(pady=15)
        
        tk.Label(dialog, text="密码:", bg='#f0f0f0').pack(pady=2)
        pwd1 = tk.Entry(dialog, show="*", width=30, font=("微软雅黑", 10))
        pwd1.pack(pady=5)
        pwd1.focus_set()
        
        tk.Label(dialog, text="确认密码:", bg='#f0f0f0').pack(pady=2)
        pwd2 = tk.Entry(dialog, show="*", width=30, font=("微软雅黑", 10))
        pwd2.pack(pady=5)
        
        status_lbl = tk.Label(dialog, text="", bg='#f0f0f0', fg='red')
        status_lbl.pack()

        def on_encrypt_click():
            pw1 = pwd1.get()
            pw2 = pwd2.get()
            
            if not pw1:
                status_lbl.config(text="请输入密码")
                return
            if pw1 != pw2:
                status_lbl.config(text="两次密码不一致")
                return
            
            dialog.destroy()
            self.encrypt_and_save_action(pw1)

        # 明显的加密按钮
        btn_frame = tk.Frame(dialog, bg='#f0f0f0')
        btn_frame.pack(pady=10)
        tk.Button(btn_frame, text="🚀 立即加密并保存", command=on_encrypt_click, 
                  bg='#e74c3c', fg='white', font=("微软雅黑", 11, "bold"), padx=30, pady=8, relief=tk.RAISED, bd=3).pack()

    def encrypt_and_save_action(self, password):
        """执行保存并应用页边距（修复加密API参数错误）"""
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("文档 (加密)", "*.docx")],
            title="保存加密文档"
        )
        if not file_path:
            return

        try:
            # 1. 生成 DOCX 并应用页边距
            in_memory_stream = self._generate_docx_with_margins()
            
            # 2. 加密（修复：直接在encrypt方法中传入password和outfile参数）
            with open(file_path, "wb") as f:
                office_file = msoffcrypto.OfficeFile(in_memory_stream)
                office_file.encrypt(password=password, outfile=f)  # 关键修复：合并参数
            
            self.is_modified = False
            self.current_file = file_path
            self.update_title()
            messagebox.showinfo("成功", "文档已加密保存！")
            
        except Exception as e:
            messagebox.showerror("错误", f"保存失败：{str(e)}")

    def _generate_docx_with_margins(self):
        """辅助函数：生成带格式和边距的docx内存文件"""
        content = self.text_area.get(1.0, tk.END).strip()
        doc = docx.Document()
        
        # --- 应用页边距 ---
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(self.page_margins['top'])
            section.bottom_margin = Cm(self.page_margins['bottom'])
            section.left_margin = Cm(self.page_margins['left'])
            section.right_margin = Cm(self.page_margins['right'])
        
        # --- 写入内容 ---
        for line in content.split('\n'):
            if line.strip():
                para = doc.add_paragraph()
                if line.startswith("<center>"):
                    line = line[8:-9]
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif line.startswith("<right>"):
                    line = line[7:-8]
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                para.add_run(line)
        
        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)
        return stream

    # ================= 修复缺失的方法：项目符号 + 编号 =================
    def add_bullets(self):
        """添加项目符号"""
        try:
            start = self.text_area.index("insert linestart")
            self.text_area.insert(start, "• ")
            self.is_modified = True
        except: pass

    def add_numbering(self):
        """添加编号"""
        try:
            # 获取当前段落行数
            line_num = self.text_area.index("insert").split('.')[0]
            start = self.text_area.index("insert linestart")
            self.text_area.insert(start, f"{line_num}. ")
            self.is_modified = True
        except: pass

    # ================= 通用功能 =================
    
    def new_file(self):
        if self.check_save():
            self.text_area.config(state='normal')
            self.text_area.delete(1.0, tk.END)
            self.current_file = None
            self.is_modified = False
            self.update_title()
    
    def open_file(self):
        if not self.check_save(): return
        file_path = filedialog.askopenfilename(
            title="打开文件", filetypes=[("文档", "*.docx"), ("文本文件", "*.txt")]
        )
        if file_path:
            try:
                self.text_area.config(state='normal')
                self.text_area.delete(1.0, tk.END)
                if file_path.endswith('.docx'):
                    doc = docx.Document(file_path)
                    for para in doc.paragraphs:
                        text = para.text
                        if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                            self.text_area.insert(tk.END, text + "\n", "center")
                        else:
                            self.text_area.insert(tk.END, text + "\n")
                else:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        self.text_area.insert(1.0, f.read())
                self.current_file = file_path
                self.is_modified = False
                self.update_title()
            except Exception as e:
                messagebox.showerror("错误", f"打开失败：{str(e)}")
    
    def save_file(self):
        if self.current_file:
            if self.current_file.endswith('.txt'):
                with open(self.current_file, 'w', encoding='utf-8') as f:
                    f.write(self.text_area.get(1.0, tk.END))
                self.is_modified = False
                self.update_title()
                return True
            else:
                return self.save_as_file()
        else:
            return self.save_as_file()
    
    def save_as_file(self):
        path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("文档", "*.docx"), ("文本文件", "*.txt")])
        if path:
            if path.endswith('.docx'):
                # 普通保存也应用页边距
                stream = self._generate_docx_with_margins()
                with open(path, 'wb') as f:
                    f.write(stream.getvalue())
            else:
                with open(path, 'w', encoding='utf-8') as f:
                    f.write(self.text_area.get(1.0, tk.END))
            self.current_file = path
            self.is_modified = False
            self.update_title()
            return True
        return False

    def check_save(self):
        """关闭前的检查提示"""
        if self.is_modified:
            result = messagebox.askyesnocancel("保存提示", "当前文档已修改，是否保存更改？")
            if result:  # 是
                return self.save_file()
            elif result is None: # 取消
                return False
        return True # 没修改或者选了“否”

    def on_closing(self):
        """窗口关闭事件"""
        if self.check_save():
            self.root.destroy()

    def change_font_family(self, event=None):
        self.text_area.config(font=(self.font_family_combo.get(), int(self.font_size_combo.get())))
    
    def change_font_size(self, event=None):
        self.text_area.config(font=(self.font_family_combo.get(), int(self.font_size_combo.get())))
    
    def bold_text(self):
        try:
            if self.text_area.tag_ranges(tk.SEL):
                self.text_area.tag_add("bold", tk.SEL_FIRST, tk.SEL_LAST)
        except: pass
    
    def italic_text(self):
        try:
            if self.text_area.tag_ranges(tk.SEL):
                self.text_area.tag_add("italic", tk.SEL_FIRST, tk.SEL_LAST)
        except: pass
    
    def underline_text(self):
        try:
            if self.text_area.tag_ranges(tk.SEL):
                self.text_area.tag_add("underline", tk.SEL_FIRST, tk.SEL_LAST)
        except: pass
    
    def choose_text_color(self):
        color = colorchooser.askcolor()[1]
        if color: self.text_area.config(foreground=color)
    
    def set_alignment(self, align):
        tag = 'left' if align == 'left' else ('center' if align == 'center' else 'right')
        try:
            start = self.text_area.index("insert linestart")
            end = self.text_area.index("insert lineend")
            self.text_area.tag_add(tag, start, end)
            self.text_area.tag_config(tag, justify=align)
        except: pass
    
    def insert_date(self): self.text_area.insert(tk.INSERT, datetime.now().strftime("%Y-%m-%d"))
    def insert_time(self): self.text_area.insert(tk.INSERT, datetime.now().strftime("%H:%M:%S"))
    def insert_separator(self): self.text_area.insert(tk.INSERT, "\n" + "="*60 + "\n")
    
    def insert_table(self):
        dialog = tk.Toplevel(self.root)
        dialog.title("插入表格")
        dialog.geometry("300x200")
        tk.Label(dialog, text="行数:").pack(pady=5)
        r = tk.Spinbox(dialog, from_=1, to=20)
        r.pack()
        tk.Label(dialog, text="列数:").pack(pady=5)
        c = tk.Spinbox(dialog, from_=1, to=20)
        c.pack()
        def create():
            table = "\n"
            table += "+" + "---+" * int(c.get()) + "\n"
            for _ in range(int(r.get())):
                table += "|" + "   |" * int(c.get()) + "\n"
                table += "+" + "---+" * int(c.get()) + "\n"
            self.text_area.insert(tk.INSERT, table)
            dialog.destroy()
        tk.Button(dialog, text="创建", command=create, bg='#3498db', fg='white').pack(pady=10)
    
    def undo(self): 
        try: self.text_area.edit_undo()
        except: pass
    def redo(self): 
        try: self.text_area.edit_redo()
        except: pass
    def copy_text(self): 
        try: pyperclip.copy(self.text_area.get(tk.SEL_FIRST, tk.SEL_LAST))
        except: pass
    def paste_text(self): 
        try: self.text_area.insert(tk.INSERT, pyperclip.paste())
        except: pass
    def cut_text(self):
        try: 
            pyperclip.copy(self.text_area.get(tk.SEL_FIRST, tk.SEL_LAST))
            self.text_area.delete(tk.SEL_FIRST, tk.SEL_LAST)
        except: pass
    def select_all(self):
        self.text_area.tag_add(tk.SEL, "1.0", tk.END)
    
    def find_replace(self):
        dialog = tk.Toplevel(self.root)
        dialog.title("查找替换")
        dialog.geometry("400x200")
        tk.Label(dialog, text="查找:").grid(row=0, column=0, padx=10, pady=10)
        f = tk.Entry(dialog, width=30)
        f.grid(row=0, column=1, padx=10, pady=10)
        tk.Label(dialog, text="替换:").grid(row=1, column=0, padx=10, pady=10)
        r = tk.Entry(dialog, width=30)
        r.grid(row=1, column=1, padx=10, pady=10)
        
        def rp():
            c = self.text_area.get(1.0, tk.END)
            nc = c.replace(f.get(), r.get())
            self.text_area.delete(1.0, tk.END)
            self.text_area.insert(1.0, nc)
        tk.Button(dialog, text="全部替换", command=rp, bg='#e74c3c', fg='white').grid(row=2, column=1, pady=10)
    
    def update_stats(self, event=None):
        content = self.text_area.get(1.0, tk.END)
        self.word_count_label.config(text=f"字数: {len(content.split())}")
        self.char_count_label.config(text=f"字符: {len(content.replace(chr(10), ''))}")
        self.para_count_label.config(text=f"段落: {len(content.split(chr(10)))}")
    
    def on_text_modified(self, event=None):
        self.is_modified = True
        self.update_title()
        self.text_area.edit_modified(False)
        self.update_stats()
    
    def update_title(self):
        title = "专业便携式文档编辑器"
        if self.current_file: title = f"{os.path.basename(self.current_file)} - {title}"
        if self.is_modified: title = f"* {title}"
        self.root.title(title)

if __name__ == "__main__":
    app = ModernWordEditor()
    app.root.mainloop()