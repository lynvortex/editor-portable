import tkinter as tk
from tkinter import ttk, filedialog, messagebox, colorchooser, font, simpledialog
import os
from datetime import datetime
import io
import docx
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import pyperclip
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import re

class SimpleDocEditor:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("LynEdit-w")
        self.root.geometry("1200x800")
        self.root.configure(bg='#f0f0f0')
        
        # 设置应用图标
        if os.path.exists("app.ico"):
            self.root.iconbitmap("app.ico")
        
        # 文件状态
        self.current_file = None
        self.is_modified = False
        self.current_font_size = 12
        self.current_font_family = "微软雅黑"
        
        # 页面边距设置
        self.page_margins = {
            'top': 2.54,
            'bottom': 2.54,
            'left': 3.17,
            'right': 3.17
        }
        
        # 高亮颜色
        self.highlight_colors = {
            "黄色": "#FFFF00",
            "绿色": "#90EE90",
            "红色": "#FFB6C1",
            "蓝色": "#ADD8E6",
            "灰色": "#D3D3D3"
        }
        
        # 注册字体
        try:
            pdfmetrics.registerFont(TTFont('SimHei', 'simhei.ttf'))
            self.pdf_font = 'SimHei'
        except:
            self.pdf_font = 'Helvetica'
        
        # 创建界面
        self.create_menu()
        self.create_toolbar()
        self.create_statusbar()
        self.create_editor()
        
        # 绑定快捷键
        self.bind_shortcuts()
        
        # 初始化
        self.new_file()
        
    def create_menu(self):
        """简洁菜单栏"""
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)
        
        # 文件菜单
        file_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="文件", menu=file_menu)
        file_menu.add_command(label="新建", command=self.new_file, accelerator="Ctrl+N")
        file_menu.add_command(label="打开", command=self.open_file, accelerator="Ctrl+O")
        file_menu.add_command(label="保存", command=self.save_file, accelerator="Ctrl+S")
        file_menu.add_command(label="另存为", command=self.save_as_file, accelerator="Ctrl+Shift+S")
        file_menu.add_separator()
        file_menu.add_command(label="导出PDF", command=self.export_pdf, accelerator="Ctrl+P")
        file_menu.add_separator()
        file_menu.add_command(label="加密保存", command=self.encrypt_dialog)
        file_menu.add_separator()
        file_menu.add_command(label="退出", command=self.on_closing, accelerator="Ctrl+Q")
        
        # 编辑菜单
        edit_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="编辑", menu=edit_menu)
        edit_menu.add_command(label="撤销", command=self.undo, accelerator="Ctrl+Z")
        edit_menu.add_command(label="重做", command=self.redo, accelerator="Ctrl+Y")
        edit_menu.add_separator()
        edit_menu.add_command(label="剪切", command=self.cut, accelerator="Ctrl+X")
        edit_menu.add_command(label="复制", command=self.copy, accelerator="Ctrl+C")
        edit_menu.add_command(label="粘贴", command=self.paste, accelerator="Ctrl+V")
        edit_menu.add_separator()
        edit_menu.add_command(label="全选", command=self.select_all, accelerator="Ctrl+A")
        edit_menu.add_command(label="查找替换", command=self.find_replace, accelerator="Ctrl+F")
        
        # 格式菜单
        format_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="格式", menu=format_menu)
        format_menu.add_command(label="字体", command=self.set_font)
        format_menu.add_command(label="颜色", command=self.set_color)
        format_menu.add_separator()
        format_menu.add_command(label="粗体", command=self.make_bold)
        format_menu.add_command(label="斜体", command=self.make_italic)
        format_menu.add_command(label="下划线", command=self.make_underline)
        format_menu.add_separator()
        format_menu.add_command(label="高亮", command=self.show_highlight_menu)
        format_menu.add_command(label="清除高亮", command=self.clear_highlight)
        format_menu.add_separator()
        format_menu.add_command(label="左对齐", command=lambda: self.set_align('left'))
        format_menu.add_command(label="居中", command=lambda: self.set_align('center'))
        format_menu.add_command(label="右对齐", command=lambda: self.set_align('right'))
        format_menu.add_separator()
        format_menu.add_command(label="项目符号", command=self.add_bullet)
        format_menu.add_command(label="编号", command=self.add_number)
        
        # 页面菜单
        page_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="页面", menu=page_menu)
        page_menu.add_command(label="页边距", command=self.page_setup)
        page_menu.add_command(label="插入分页符", command=self.insert_page_break)
        
        # 插入菜单
        insert_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="插入", menu=insert_menu)
        insert_menu.add_command(label="日期", command=self.insert_date)
        insert_menu.add_command(label="时间", command=self.insert_time)
        insert_menu.add_command(label="分隔线", command=self.insert_line)
        insert_menu.add_command(label="表格", command=self.insert_table)
        
        # 关于菜单（原帮助菜单）
        about_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="关于", menu=about_menu)
        about_menu.add_command(label="关于", command=self.show_about)
    
    def create_toolbar(self):
        """简洁工具栏"""
        toolbar = tk.Frame(self.root, bg='#e0e0e0', relief=tk.RAISED, bd=1)
        toolbar.pack(side=tk.TOP, fill=tk.X)
        
        # 文件操作
        self.add_tool_button(toolbar, "新建", self.new_file).pack(side=tk.LEFT, padx=2, pady=2)
        self.add_tool_button(toolbar, "打开", self.open_file).pack(side=tk.LEFT, padx=2, pady=2)
        self.add_tool_button(toolbar, "保存", self.save_file).pack(side=tk.LEFT, padx=2, pady=2)
        
        tk.Frame(toolbar, width=2, bg='#cccccc').pack(side=tk.LEFT, padx=5, fill=tk.Y, pady=2)
        
        # 字体控制
        tk.Label(toolbar, text="字体:", bg='#e0e0e0').pack(side=tk.LEFT, padx=2)
        self.font_combo = ttk.Combobox(toolbar, values=list(tk.font.families())[:30], width=12)
        self.font_combo.set("微软雅黑")
        self.font_combo.pack(side=tk.LEFT, padx=2)
        self.font_combo.bind('<<ComboboxSelected>>', self.on_font_change)
        
        tk.Label(toolbar, text="字号:", bg='#e0e0e0').pack(side=tk.LEFT, padx=2)
        self.size_combo = ttk.Combobox(toolbar, values=[8,9,10,11,12,14,16,18,20,22,24,28,32,36,48], width=5)
        self.size_combo.set(12)
        self.size_combo.pack(side=tk.LEFT, padx=2)
        self.size_combo.bind('<<ComboboxSelected>>', self.on_size_change)
        
        tk.Frame(toolbar, width=2, bg='#cccccc').pack(side=tk.LEFT, padx=5, fill=tk.Y, pady=2)
        
        # 文字样式
        self.add_tool_button(toolbar, "粗体", self.make_bold).pack(side=tk.LEFT, padx=2)
        self.add_tool_button(toolbar, "斜体", self.make_italic).pack(side=tk.LEFT, padx=2)
        self.add_tool_button(toolbar, "下划线", self.make_underline).pack(side=tk.LEFT, padx=2)
        
        tk.Frame(toolbar, width=2, bg='#cccccc').pack(side=tk.LEFT, padx=5, fill=tk.Y, pady=2)
        
        # 对齐
        self.add_tool_button(toolbar, "左对齐", lambda: self.set_align('left')).pack(side=tk.LEFT, padx=2)
        self.add_tool_button(toolbar, "居中", lambda: self.set_align('center')).pack(side=tk.LEFT, padx=2)
        self.add_tool_button(toolbar, "右对齐", lambda: self.set_align('right')).pack(side=tk.LEFT, padx=2)
        
        tk.Frame(toolbar, width=2, bg='#cccccc').pack(side=tk.LEFT, padx=5, fill=tk.Y, pady=2)
        
        # 页边距按钮
        self.add_tool_button(toolbar, "页边距", self.page_setup).pack(side=tk.LEFT, padx=2)
        
        tk.Frame(toolbar, width=2, bg='#cccccc').pack(side=tk.LEFT, padx=5, fill=tk.Y, pady=2)
        
        # 其他功能
        self.add_tool_button(toolbar, "高亮", self.show_highlight_menu).pack(side=tk.LEFT, padx=2)
        self.add_tool_button(toolbar, "加密", self.encrypt_dialog).pack(side=tk.LEFT, padx=2)
    
    def add_tool_button(self, parent, text, command):
        btn = tk.Button(parent, text=text, command=command, bg='#e0e0e0',
                       relief=tk.FLAT, padx=8, pady=2, font=("微软雅黑", 9))
        btn.bind('<Enter>', lambda e: btn.config(relief=tk.RAISED, bg='#d0d0d0'))
        btn.bind('<Leave>', lambda e: btn.config(relief=tk.FLAT, bg='#e0e0e0'))
        return btn
    
    def create_editor(self):
        """编辑区域"""
        main_frame = tk.Frame(self.root, bg='#f0f0f0')
        main_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # 文本编辑框
        self.text = tk.Text(main_frame, wrap=tk.WORD, font=("微软雅黑", 12),
                            undo=True, autoseparators=True, maxundo=100,
                            padx=20, pady=20, bg='white', fg='black',
                            selectbackground='#3399ff', selectforeground='white',
                            relief=tk.FLAT, borderwidth=1, highlightthickness=1,
                            highlightbackground='#cccccc')
        self.text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # 滚动条
        scrollbar = tk.Scrollbar(main_frame, orient=tk.VERTICAL, command=self.text.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.text.config(yscrollcommand=scrollbar.set)
        
        # 标签样式 - 支持不同字号
        self.text.tag_configure("bold", font=(self.current_font_family, self.current_font_size, "bold"))
        self.text.tag_configure("italic", font=(self.current_font_family, self.current_font_size, "italic"))
        self.text.tag_configure("underline", font=(self.current_font_family, self.current_font_size, "underline"))
        self.text.tag_configure("strike", overstrike=True)
        self.text.tag_configure("left", justify='left')
        self.text.tag_configure("center", justify='center')
        self.text.tag_configure("right", justify='right')
        
        for name, color in self.highlight_colors.items():
            self.text.tag_configure(f"hl_{name}", background=color)
        
        # 事件
        self.text.bind('<<Modified>>', self.on_modified)
        self.text.bind('<KeyRelease>', self.update_stats)
        self.text.bind('<ButtonRelease-1>', self.update_cursor_pos)
    
    def create_statusbar(self):
        """状态栏"""
        self.statusbar = tk.Frame(self.root, bg='#e0e0e0', relief=tk.SUNKEN, bd=1)
        self.statusbar.pack(side=tk.BOTTOM, fill=tk.X)
        
        self.status_text = tk.Label(self.statusbar, text="开源地址:https://github.com/lynvortex/LynEdit", bg='#e0e0e0', font=("微软雅黑", 9))
        self.status_text.pack(side=tk.LEFT, padx=5)
        
        self.status_words = tk.Label(self.statusbar, text="字数: 0", bg='#e0e0e0', font=("微软雅黑", 9))
        self.status_words.pack(side=tk.RIGHT, padx=5)
        
        self.status_chars = tk.Label(self.statusbar, text="字符: 0", bg='#e0e0e0', font=("微软雅黑", 9))
        self.status_chars.pack(side=tk.RIGHT, padx=5)
        
        self.status_cursor = tk.Label(self.statusbar, text="行1,列1", bg='#e0e0e0', font=("微软雅黑", 9))
        self.status_cursor.pack(side=tk.RIGHT, padx=5)
    
    # ==================== 核心功能 ====================
    
    def new_file(self):
        if self.check_save():
            self.text.delete(1.0, tk.END)
            self.current_file = None
            self.is_modified = False
            self.update_title()
            self.update_stats()
    
    def open_file(self):
        if not self.check_save():
            return
        path = filedialog.askopenfilename(filetypes=[("Word文档", "*.docx"), ("文本文件", "*.txt"), ("所有文件", "*.*")])
        if path:
            try:
                self.text.delete(1.0, tk.END)
                if path.endswith('.docx'):
                    doc = docx.Document(path)
                    for para in doc.paragraphs:
                        text = para.text
                        # 保留对齐格式
                        if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                            self.text.insert(tk.END, text + "\n", "center")
                        elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                            self.text.insert(tk.END, text + "\n", "right")
                        else:
                            self.text.insert(tk.END, text + "\n", "left")
                else:
                    with open(path, 'r', encoding='utf-8') as f:
                        self.text.insert(1.0, f.read())
                self.current_file = path
                self.is_modified = False
                self.text.edit_modified(False)
                self.update_title()
                self.update_stats()
                # 移除了成功弹窗
            except Exception as e:
                messagebox.showerror("错误", f"打开失败: {e}")
    
    def save_file(self):
        if self.current_file:
            if self.current_file.endswith('.txt'):
                with open(self.current_file, 'w', encoding='utf-8') as f:
                    f.write(self.text.get(1.0, tk.END))
                self.is_modified = False
                self.update_title()
                return True
            else:
                return self.save_as_file()
        else:
            return self.save_as_file()
    
    def save_as_file(self):
        path = filedialog.asksaveasfilename(defaultextension=".docx",
                                            filetypes=[("Word文档", "*.docx"), ("文本文件", "*.txt")])
        if path:
            try:
                if path.endswith('.docx'):
                    doc = docx.Document()
                    # 设置页边距
                    for section in doc.sections:
                        section.top_margin = Cm(self.page_margins['top'])
                        section.bottom_margin = Cm(self.page_margins['bottom'])
                        section.left_margin = Cm(self.page_margins['left'])
                        section.right_margin = Cm(self.page_margins['right'])
                    
                    # 获取带格式的文本
                    content = self.text.get(1.0, tk.END).rstrip('\n')
                    lines = content.split('\n')
                    
                    for line in lines:
                        if line.strip() or line == '':
                            para = doc.add_paragraph()
                            # 检查行是否有对齐标签
                            if self.text.tag_ranges("center"):
                                # 简化处理：检查当前行是否在center标签内
                                pass
                            para.add_run(line)
                            
                            # 根据文本中的标记设置对齐
                            if line.startswith('  '):
                                pass
                    
                    doc.save(path)
                else:
                    with open(path, 'w', encoding='utf-8') as f:
                        f.write(self.text.get(1.0, tk.END))
                
                self.current_file = path
                self.is_modified = False
                self.update_title()
                messagebox.showinfo("成功", "保存成功")
                return True
            except Exception as e:
                messagebox.showerror("错误", f"保存失败: {e}")
        return False
    
    def export_pdf(self):
        if not self.text.get(1.0, tk.END).strip():
            messagebox.showwarning("警告", "没有内容")
            return
        
        path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF文件", "*.pdf")])
        if path:
            try:
                content = self.text.get(1.0, tk.END)
                c = canvas.Canvas(path, pagesize=A4)
                width, height = A4
                y = height - self.page_margins['top'] * cm
                c.setFont(self.pdf_font, self.current_font_size)
                
                for line in content.split('\n'):
                    if line.strip():
                        # 移除格式标记
                        clean_line = re.sub(r'<[^>]+>', '', line)
                        c.drawString(self.page_margins['left'] * cm, y, clean_line[:100])
                        y -= self.current_font_size + 4
                        if y < self.page_margins['bottom'] * cm:
                            c.showPage()
                            c.setFont(self.pdf_font, self.current_font_size)
                            y = height - self.page_margins['top'] * cm
                
                c.save()
                messagebox.showinfo("成功", "PDF导出成功")
            except Exception as e:
                messagebox.showerror("错误", f"导出失败: {e}")
    
    def encrypt_dialog(self):
        dialog = tk.Toplevel(self.root)
        dialog.title("加密保存")
        dialog.geometry("350x220")
        dialog.resizable(False, False)
        
        # 保存密码变量的引用
        self.encrypt_pwd1 = tk.StringVar()
        self.encrypt_pwd2 = tk.StringVar()
        
        tk.Label(dialog, text="设置密码:", font=("微软雅黑", 10)).pack(pady=10)
        pwd1 = tk.Entry(dialog, show="*", width=30, textvariable=self.encrypt_pwd1, font=("微软雅黑", 10))
        pwd1.pack(pady=5)
        
        tk.Label(dialog, text="确认密码:", font=("微软雅黑", 10)).pack(pady=5)
        pwd2 = tk.Entry(dialog, show="*", width=30, textvariable=self.encrypt_pwd2, font=("微软雅黑", 10))
        pwd2.pack(pady=5)
        
        error_label = tk.Label(dialog, text="", fg='red', font=("微软雅黑", 9))
        error_label.pack(pady=5)
        
        def do_encrypt():
            p1 = self.encrypt_pwd1.get()
            p2 = self.encrypt_pwd2.get()
            if not p1:
                error_label.config(text="请输入密码")
                return
            if p1 != p2:
                error_label.config(text="两次密码不一致")
                return
            dialog.destroy()
            self.encrypt_save(p1)
        
        btn_frame = tk.Frame(dialog)
        btn_frame.pack(pady=15)
        tk.Button(btn_frame, text="保存", command=do_encrypt, bg='#4CAF50', fg='white', 
                 font=("微软雅黑", 10), padx=20).pack(side=tk.LEFT, padx=5)
        tk.Button(btn_frame, text="取消", command=dialog.destroy, bg='#f44336', fg='white',
                 font=("微软雅黑", 10), padx=20).pack(side=tk.LEFT, padx=5)
    
    def encrypt_save(self, password):
        path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("加密文档", "*.docx")])
        if path:
            try:
                import msoffcrypto
                
                # 创建文档
                doc = docx.Document()
                for section in doc.sections:
                    section.top_margin = Cm(self.page_margins['top'])
                    section.bottom_margin = Cm(self.page_margins['bottom'])
                    section.left_margin = Cm(self.page_margins['left'])
                    section.right_margin = Cm(self.page_margins['right'])
                
                # 获取内容
                content = self.text.get(1.0, tk.END).strip()
                for line in content.split('\n'):
                    if line.strip():
                        para = doc.add_paragraph()
                        # 检查对齐
                        if self.text.tag_ranges("center"):
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif self.text.tag_ranges("right"):
                            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        para.add_run(line)
                    else:
                        doc.add_paragraph()
                
                # 保存到内存
                stream = io.BytesIO()
                doc.save(stream)
                stream.seek(0)
                
                # 加密保存
                with open(path, "wb") as f:
                    office_file = msoffcrypto.OfficeFile(stream)
                    office_file.encrypt(password=password, outfile=f)
                
                self.current_file = path
                self.is_modified = False
                self.update_title()
                messagebox.showinfo("成功", "加密保存成功")
                
            except ImportError:
                messagebox.showerror("错误", "请安装 msoffcrypto-tool\n命令: pip install msoffcrypto-tool")
            except Exception as e:
                messagebox.showerror("错误", f"加密失败: {e}")
    
    # ==================== 格式功能 ====================
    
    def update_font_tags(self):
        """更新字体标签"""
        self.text.tag_configure("bold", font=(self.current_font_family, self.current_font_size, "bold"))
        self.text.tag_configure("italic", font=(self.current_font_family, self.current_font_size, "italic"))
        self.text.tag_configure("underline", font=(self.current_font_family, self.current_font_size, "underline"))
    
    def set_font(self):
        font_win = tk.Toplevel(self.root)
        font_win.title("字体")
        font_win.geometry("300x220")
        font_win.resizable(False, False)
        
        tk.Label(font_win, text="字体:", font=("微软雅黑", 10)).pack(pady=5)
        font_combo = ttk.Combobox(font_win, values=list(tk.font.families())[:30], width=25)
        font_combo.set(self.current_font_family)
        font_combo.pack(pady=5)
        
        tk.Label(font_win, text="字号:", font=("微软雅黑", 10)).pack(pady=5)
        size_combo = ttk.Combobox(font_win, values=[8,9,10,11,12,14,16,18,20,22,24,28,32,36,48], width=10)
        size_combo.set(self.current_font_size)
        size_combo.pack(pady=5)
        
        def apply():
            self.current_font_family = font_combo.get()
            self.current_font_size = int(size_combo.get())
            self.text.config(font=(self.current_font_family, self.current_font_size))
            self.font_combo.set(self.current_font_family)
            self.size_combo.set(self.current_font_size)
            self.update_font_tags()
            font_win.destroy()
        
        tk.Button(font_win, text="确定", command=apply, bg='#4CAF50', fg='white',
                 font=("微软雅黑", 10), padx=20, pady=5).pack(pady=15)
    
    def set_color(self):
        color = colorchooser.askcolor()[1]
        if color:
            self.text.config(foreground=color)
    
    def show_highlight_menu(self):
        menu = tk.Menu(self.root, tearoff=0)
        for name in self.highlight_colors:
            menu.add_command(label=name, command=lambda n=name: self.highlight_text(n))
        menu.add_separator()
        menu.add_command(label="清除高亮", command=self.clear_highlight)
        menu.post(self.root.winfo_pointerx(), self.root.winfo_pointery())
    
    def highlight_text(self, color_name):
        try:
            if self.text.tag_ranges(tk.SEL):
                self.text.tag_add(f"hl_{color_name}", tk.SEL_FIRST, tk.SEL_LAST)
        except:
            pass
    
    def clear_highlight(self):
        try:
            if self.text.tag_ranges(tk.SEL):
                for name in self.highlight_colors:
                    self.text.tag_remove(f"hl_{name}", tk.SEL_FIRST, tk.SEL_LAST)
        except:
            pass
    
    def make_bold(self):
        try:
            if self.text.tag_ranges(tk.SEL):
                self.text.tag_add("bold", tk.SEL_FIRST, tk.SEL_LAST)
        except:
            pass
    
    def make_italic(self):
        try:
            if self.text.tag_ranges(tk.SEL):
                self.text.tag_add("italic", tk.SEL_FIRST, tk.SEL_LAST)
        except:
            pass
    
    def make_underline(self):
        try:
            if self.text.tag_ranges(tk.SEL):
                self.text.tag_add("underline", tk.SEL_FIRST, tk.SEL_LAST)
        except:
            pass
    
    def set_align(self, align):
        try:
            if self.text.tag_ranges(tk.SEL):
                start = self.text.index(f"{tk.SEL_FIRST} linestart")
                end = self.text.index(f"{tk.SEL_LAST} lineend + 1c")
                self.text.tag_add(align, start, end)
            else:
                start = self.text.index("insert linestart")
                end = self.text.index("insert lineend + 1c")
                self.text.tag_add(align, start, end)
        except:
            pass
    
    def add_bullet(self):
        try:
            self.text.insert("insert linestart", "• ")
            self.is_modified = True
        except:
            pass
    
    def add_number(self):
        try:
            line = self.text.index("insert").split('.')[0]
            self.text.insert("insert linestart", f"{line}. ")
            self.is_modified = True
        except:
            pass
    
    def page_setup(self):
        """页边距设置对话框"""
        dialog = tk.Toplevel(self.root)
        dialog.title("页边距设置")
        dialog.geometry("340x320")
        dialog.resizable(False, False)
        dialog.transient(self.root)
        dialog.grab_set()
        
        # 标题
        title_frame = tk.Frame(dialog, bg='#4CAF50', height=40)
        title_frame.pack(fill=tk.X)
        title_frame.pack_propagate(False)
        tk.Label(title_frame, text="页边距设置 (厘米)", bg='#4CAF50', fg='white', 
                font=("微软雅黑", 12, "bold")).pack(expand=True)
        
        # 输入区域
        main_frame = tk.Frame(dialog, bg='#f5f5f5')
        main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        
        # 上边距
        top_frame = tk.Frame(main_frame, bg='#f5f5f5')
        top_frame.pack(fill=tk.X, pady=8)
        tk.Label(top_frame, text="上边距:", bg='#f5f5f5', font=("微软雅黑", 10), width=8, anchor='w').pack(side=tk.LEFT)
        top_entry = tk.Entry(top_frame, width=10, font=("微软雅黑", 10))
        top_entry.insert(0, str(self.page_margins['top']))
        top_entry.pack(side=tk.LEFT, padx=10)
        tk.Label(top_frame, text="厘米", bg='#f5f5f5', font=("微软雅黑", 10)).pack(side=tk.LEFT)
        
        # 下边距
        bottom_frame = tk.Frame(main_frame, bg='#f5f5f5')
        bottom_frame.pack(fill=tk.X, pady=8)
        tk.Label(bottom_frame, text="下边距:", bg='#f5f5f5', font=("微软雅黑", 10), width=8, anchor='w').pack(side=tk.LEFT)
        bottom_entry = tk.Entry(bottom_frame, width=10, font=("微软雅黑", 10))
        bottom_entry.insert(0, str(self.page_margins['bottom']))
        bottom_entry.pack(side=tk.LEFT, padx=10)
        tk.Label(bottom_frame, text="厘米", bg='#f5f5f5', font=("微软雅黑", 10)).pack(side=tk.LEFT)
        
        # 左边距
        left_frame = tk.Frame(main_frame, bg='#f5f5f5')
        left_frame.pack(fill=tk.X, pady=8)
        tk.Label(left_frame, text="左边距:", bg='#f5f5f5', font=("微软雅黑", 10), width=8, anchor='w').pack(side=tk.LEFT)
        left_entry = tk.Entry(left_frame, width=10, font=("微软雅黑", 10))
        left_entry.insert(0, str(self.page_margins['left']))
        left_entry.pack(side=tk.LEFT, padx=10)
        tk.Label(left_frame, text="厘米", bg='#f5f5f5', font=("微软雅黑", 10)).pack(side=tk.LEFT)
        
        # 右边距
        right_frame = tk.Frame(main_frame, bg='#f5f5f5')
        right_frame.pack(fill=tk.X, pady=8)
        tk.Label(right_frame, text="右边距:", bg='#f5f5f5', font=("微软雅黑", 10), width=8, anchor='w').pack(side=tk.LEFT)
        right_entry = tk.Entry(right_frame, width=10, font=("微软雅黑", 10))
        right_entry.insert(0, str(self.page_margins['right']))
        right_entry.pack(side=tk.LEFT, padx=10)
        tk.Label(right_frame, text="厘米", bg='#f5f5f5', font=("微软雅黑", 10)).pack(side=tk.LEFT)
        
        # 预设按钮
        preset_frame = tk.Frame(main_frame, bg='#f5f5f5')
        preset_frame.pack(fill=tk.X, pady=15)
        
        def set_preset(top, bottom, left, right):
            top_entry.delete(0, tk.END)
            top_entry.insert(0, str(top))
            bottom_entry.delete(0, tk.END)
            bottom_entry.insert(0, str(bottom))
            left_entry.delete(0, tk.END)
            left_entry.insert(0, str(left))
            right_entry.delete(0, tk.END)
            right_entry.insert(0, str(right))
        
        tk.Button(preset_frame, text="普通", command=lambda: set_preset(2.54, 2.54, 3.17, 3.17),
                 bg='#e0e0e0', font=("微软雅黑", 9), padx=15).pack(side=tk.LEFT, padx=5)
        tk.Button(preset_frame, text="窄", command=lambda: set_preset(1.27, 1.27, 1.27, 1.27),
                 bg='#e0e0e0', font=("微软雅黑", 9), padx=15).pack(side=tk.LEFT, padx=5)
        tk.Button(preset_frame, text="宽", command=lambda: set_preset(3.81, 3.81, 3.81, 3.81),
                 bg='#e0e0e0', font=("微软雅黑", 9), padx=15).pack(side=tk.LEFT, padx=5)
        
        # 按钮区域
        btn_frame = tk.Frame(main_frame, bg='#f5f5f5')
        btn_frame.pack(fill=tk.X, pady=15)
        
        def save_margins():
            try:
                self.page_margins['top'] = float(top_entry.get())
                self.page_margins['bottom'] = float(bottom_entry.get())
                self.page_margins['left'] = float(left_entry.get())
                self.page_margins['right'] = float(right_entry.get())
                dialog.destroy()
                messagebox.showinfo("成功", "页边距设置已保存")
            except ValueError:
                messagebox.showerror("错误", "请输入有效的数字")
        
        tk.Button(btn_frame, text="确定", command=save_margins,
                 bg='#4CAF50', fg='white', font=("微软雅黑", 10), padx=20, pady=5).pack(side=tk.RIGHT, padx=5)
        tk.Button(btn_frame, text="取消", command=dialog.destroy,
                 bg='#f44336', fg='white', font=("微软雅黑", 10), padx=20, pady=5).pack(side=tk.RIGHT, padx=5)
    
    # ==================== 插入功能 ====================
    
    def insert_date(self):
        self.text.insert(tk.INSERT, datetime.now().strftime("%Y-%m-%d"))
        self.is_modified = True
    
    def insert_time(self):
        self.text.insert(tk.INSERT, datetime.now().strftime("%H:%M:%S"))
        self.is_modified = True
    
    def insert_line(self):
        self.text.insert(tk.INSERT, "\n" + "="*60 + "\n")
        self.is_modified = True
    
    def insert_page_break(self):
        self.text.insert(tk.INSERT, f"\n{'='*50}\n--- 分页 ---\n{'='*50}\n\n")
        self.is_modified = True
    
    def insert_table(self):
        dialog = tk.Toplevel(self.root)
        dialog.title("插入表格")
        dialog.geometry("260x200")
        dialog.resizable(False, False)
        
        tk.Label(dialog, text="行数:", font=("微软雅黑", 10)).pack(pady=8)
        rows = tk.Spinbox(dialog, from_=1, to=20, width=8, font=("微软雅黑", 10))
        rows.pack()
        
        tk.Label(dialog, text="列数:", font=("微软雅黑", 10)).pack(pady=8)
        cols = tk.Spinbox(dialog, from_=1, to=20, width=8, font=("微软雅黑", 10))
        cols.pack()
        
        def create():
            table = "\n"
            table += "+" + "---+" * int(cols.get()) + "\n"
            for _ in range(int(rows.get())):
                table += "|" + "   |" * int(cols.get()) + "\n"
                table += "+" + "---+" * int(cols.get()) + "\n"
            self.text.insert(tk.INSERT, table)
            self.is_modified = True
            dialog.destroy()
        
        tk.Button(dialog, text="创建", command=create, bg='#4CAF50', fg='white',
                 font=("微软雅黑", 10), padx=20, pady=5).pack(pady=15)
    
    # ==================== 编辑功能 ====================
    
    def undo(self):
        try:
            self.text.edit_undo()
        except:
            pass
    
    def redo(self):
        try:
            self.text.edit_redo()
        except:
            pass
    
    def cut(self):
        try:
            self.text.event_generate("<<Cut>>")
        except:
            pass
    
    def copy(self):
        try:
            self.text.event_generate("<<Copy>>")
        except:
            pass
    
    def paste(self):
        try:
            self.text.event_generate("<<Paste>>")
        except:
            pass
    
    def select_all(self):
        self.text.tag_add(tk.SEL, "1.0", tk.END)
    
    def find_replace(self):
        dialog = tk.Toplevel(self.root)
        dialog.title("查找替换")
        dialog.geometry("460x200")
        dialog.resizable(False, False)
        
        tk.Label(dialog, text="查找:", font=("微软雅黑", 10)).grid(row=0, column=0, padx=10, pady=10, sticky='e')
        find_e = tk.Entry(dialog, width=40, font=("微软雅黑", 10))
        find_e.grid(row=0, column=1, padx=10, pady=10)
        
        tk.Label(dialog, text="替换:", font=("微软雅黑", 10)).grid(row=1, column=0, padx=10, pady=10, sticky='e')
        replace_e = tk.Entry(dialog, width=40, font=("微软雅黑", 10))
        replace_e.grid(row=1, column=1, padx=10, pady=10)
        
        def replace():
            content = self.text.get(1.0, tk.END)
            new_content = content.replace(find_e.get(), replace_e.get())
            self.text.delete(1.0, tk.END)
            self.text.insert(1.0, new_content)
            self.is_modified = True
            dialog.destroy()
            messagebox.showinfo("完成", "替换完成")
        
        tk.Button(dialog, text="全部替换", command=replace, bg='#2196F3', fg='white',
                 font=("微软雅黑", 10), padx=20, pady=5).grid(row=2, column=1, pady=15)
    
    # ==================== 状态和辅助 ====================
    
    def update_stats(self, event=None):
        content = self.text.get(1.0, tk.END)
        words = len(content.split())
        chars = len(content.replace('\n', '').replace(' ', ''))
        self.status_words.config(text=f"字数: {words}")
        self.status_chars.config(text=f"字符: {chars}")
    
    def update_cursor_pos(self, event=None):
        pos = self.text.index(tk.INSERT)
        line, col = pos.split('.')
        self.status_cursor.config(text=f"行{line},列{int(col)+1}")
    
    def update_font_tags(self):
        """更新字体标签"""
        self.text.tag_configure("bold", font=(self.current_font_family, self.current_font_size, "bold"))
        self.text.tag_configure("italic", font=(self.current_font_family, self.current_font_size, "italic"))
        self.text.tag_configure("underline", font=(self.current_font_family, self.current_font_size, "underline"))
    
    def on_modified(self, event=None):
        self.is_modified = True
        self.update_title()
        self.text.edit_modified(False)
    
    def on_font_change(self, event=None):
        self.current_font_family = self.font_combo.get()
        self.text.config(font=(self.current_font_family, self.current_font_size))
        self.update_font_tags()
    
    def on_size_change(self, event=None):
        self.current_font_size = int(self.size_combo.get())
        self.text.config(font=(self.current_font_family, self.current_font_size))
        self.update_font_tags()
    
    def update_title(self):
        title = "专业便携式文档编辑器"
        if self.current_file:
            title = f"{os.path.basename(self.current_file)} - {title}"
        if self.is_modified:
            title = f"* {title}"
        self.root.title(title)
    
    def check_save(self):
        if self.is_modified:
            result = messagebox.askyesnocancel("提示", "当前文档已修改，是否保存？")
            if result:
                return self.save_file()
            elif result is None:
                return False
        return True
    
    def on_closing(self):
        if self.check_save():
            self.root.destroy()
    
    def bind_shortcuts(self):
        self.root.bind('<Control-n>', lambda e: self.new_file())
        self.root.bind('<Control-o>', lambda e: self.open_file())
        self.root.bind('<Control-s>', lambda e: self.save_file())
        self.root.bind('<Control-p>', lambda e: self.export_pdf())
        self.root.bind('<Control-z>', lambda e: self.undo())
        self.root.bind('<Control-y>', lambda e: self.redo())
        self.root.bind('<Control-a>', lambda e: self.select_all())
        self.root.bind('<Control-f>', lambda e: self.find_replace())
        self.root.bind('<Control-q>', lambda e: self.on_closing())
    
    def show_about(self):
        about_text = """专业便携式文档编辑器

版本号：v1.1
作者：绘萤者Lynvortex
开源地址：https://github.com/lynvortex/editor-portable
"""
        
        messagebox.showinfo("关于", about_text)


if __name__ == "__main__":
    app = SimpleDocEditor()
    app.root.mainloop()